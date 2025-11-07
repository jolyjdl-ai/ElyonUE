#!/usr/bin/env python3
"""
Génère les versions PDF et DOCX de ELYON_REFERENCE_COMPLETE.md
Dépendances: pip install python-docx python-docx-template markdown2 weasyprint
"""

import sys
from pathlib import Path

def generate_pdf():
    """Génère PDF via markdown2 + weasyprint"""
    try:
        import markdown2
        from weasyprint import HTML

        md_file = Path(__file__).parent / "ELYON_REFERENCE_COMPLETE.md"
        pdf_file = md_file.with_suffix(".pdf")

        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convertir markdown en HTML
        html_content = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks'])

        # Ajouter CSS de base
        html_with_css = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; color: #333; }}
                h1 {{ page-break-before: always; border-bottom: 3px solid #0078d4; padding-bottom: 10px; }}
                h2 {{ color: #0078d4; margin-top: 20px; }}
                h3 {{ color: #106ebe; }}
                code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
                pre {{ background: #f4f4f4; padding: 10px; border-left: 4px solid #0078d4; overflow-x: auto; }}
                table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background: #f0f0f0; }}
                ul, ol {{ margin-left: 20px; }}
                blockquote {{ border-left: 4px solid #0078d4; padding-left: 15px; margin: 15px 0; color: #666; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        # Générer PDF
        HTML(string=html_with_css).write_pdf(pdf_file)
        print(f"✅ PDF généré: {pdf_file}")

    except ImportError as e:
        print(f"❌ Dépendances manquantes: {e}")
        print("Installation: pip install markdown2 weasyprint")
        return False
    except Exception as e:
        print(f"❌ Erreur PDF: {e}")
        return False

    return True


def generate_docx():
    """Génère DOCX via python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import re

        md_file = Path(__file__).parent / "ELYON_REFERENCE_COMPLETE.md"
        docx_file = md_file.with_suffix(".docx")

        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Créer document Word
        doc = Document()

        # Ajouter titre
        title = doc.add_paragraph()
        title_run = title.add_run("ÉLYÔN — Référence Complète")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 120, 212)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Ajouter sous-titre
        subtitle = doc.add_paragraph("Document de référence pour développeurs et stakeholders")
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True

        # Table des matières
        doc.add_page_break()
        toc_title = doc.add_heading("Table des Matières", level=1)
        toc_title.runs[0].font.color.rgb = RGBColor(0, 120, 212)

        # Parser et ajouter contenu
        lines = md_content.split('\n')
        for line in lines:
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
                doc.add_heading(line[3:], level=2).runs[0].font.color.rgb = RGBColor(16, 110, 190)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
                doc.add_heading(line[2:], level=1).runs[0].font.color.rgb = RGBColor(0, 120, 212)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('```'):
                continue  # Gérer séparément
            elif line.strip():
                doc.add_paragraph(line)

        # Ajouter pied de page
        section = doc.sections[0]
        footer = section.footer.paragraphs[0]
        footer.text = f"Élyôn — v1.0 | Généré: 7 novembre 2025"
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # Sauvegarder
        doc.save(docx_file)
        print(f"✅ DOCX généré: {docx_file}")

    except ImportError as e:
        print(f"❌ Dépendances manquantes: {e}")
        print("Installation: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ Erreur DOCX: {e}")
        return False

    return True


def main():
    print("🔄 Génération des formats alternatifs de ELYON_REFERENCE_COMPLETE...")
    print()

    pdf_ok = generate_pdf()
    docx_ok = generate_docx()

    print()
    if pdf_ok and docx_ok:
        print("✅ Tous les formats générés avec succès!")
        print()
        print("📁 Fichiers disponibles:")
        docs_dir = Path(__file__).parent
        print(f"  - {docs_dir / 'ELYON_REFERENCE_COMPLETE.md'} (source)")
        print(f"  - {docs_dir / 'ELYON_REFERENCE_COMPLETE.pdf'}")
        print(f"  - {docs_dir / 'ELYON_REFERENCE_COMPLETE.docx'}")
    else:
        print("⚠️  Certains formats n'ont pas pu être générés.")
        print("Vérifiez les dépendances et relancez le script.")
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())
